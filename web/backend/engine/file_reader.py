from engine.types import FileContent, RawSheet, MergeCell


def read_file(path: str) -> FileContent:
    ext = path.rsplit(".", 1)[-1].lower()
    if ext in ("xlsx", "xls"):
        return _read_excel(path)
    elif ext == "pdf":
        return _read_pdf(path)
    elif ext == "docx":
        return _read_word(path)
    else:
        raise ValueError(f"不支持的文件类型: .{ext}")


def _read_excel(path: str) -> FileContent:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    sheets = []
    file_name = path.rsplit("/", 1)[-1]
    total_rows = 0

    for ws in wb.worksheets:
        rows = []
        merge_info = []
        for row in ws.iter_rows(values_only=True):
            cells = []
            for v in row:
                if v is None:
                    cells.append(None)
                elif hasattr(v, "strftime"):
                    cells.append(v.strftime("%Y-%m-%d"))
                elif isinstance(v, float) and v == int(v):
                    cells.append(int(v))
                else:
                    cells.append(v)
            rows.append(cells)

        try:
            if ws.merged_cells:
                for mc in ws.merged_cells.ranges:
                    merge_info.append(MergeCell(
                        row=mc.min_row - 1,
                        col=mc.min_col - 1,
                        rowspan=mc.max_row - mc.min_row + 1,
                        colspan=mc.max_col - mc.min_col + 1,
                    ))
        except AttributeError:
            pass

        raw_text = "\n".join(
            "\t".join(str(c) for c in row if c is not None)
            for row in rows
        )
        total_rows += len(rows)
        sheets.append(RawSheet(name=ws.title, rows=rows, raw_text=raw_text, merge_info=merge_info))

    wb.close()
    full_text = "\n\n".join(f"=== Sheet: {s.name} ===\n{s.raw_text}" for s in sheets)
    return FileContent(file_name=file_name, file_type="excel", sheets=sheets, full_text=full_text, total_rows=total_rows)


def _read_pdf(path: str) -> FileContent:
    import pdfplumber
    file_name = path.rsplit("/", 1)[-1]
    all_rows = []
    raw_lines = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    raw_lines.append(line)
                    cells = line.split()
                    row = [c for c in cells]
                    all_rows.append(row)

    raw_text = "\n".join(raw_lines)
    total_rows = len(all_rows)
    sheets = [RawSheet(name="content", rows=all_rows, raw_text=raw_text)]
    full_text = f"=== Sheet: content ===\n{raw_text}"
    return FileContent(file_name=file_name, file_type="pdf", sheets=sheets, full_text=full_text, total_rows=total_rows)


def _read_word(path: str) -> FileContent:
    from docx import Document
    file_name = path.rsplit("/", 1)[-1]
    doc = Document(path)
    raw_lines = []
    rows = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            raw_lines.append(text)
            rows.append([text])

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(c for c in cells):
                raw_lines.append("\t".join(cells))
                rows.append(cells)

    raw_text = "\n".join(raw_lines)
    total_rows = len(rows)
    sheets = [RawSheet(name="content", rows=rows, raw_text=raw_text)]
    full_text = f"=== Sheet: content ===\n{raw_text}"
    return FileContent(file_name=file_name, file_type="word", sheets=sheets, full_text=full_text, total_rows=total_rows)
